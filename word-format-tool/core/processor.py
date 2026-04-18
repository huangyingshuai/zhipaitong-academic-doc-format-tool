from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO

def is_protected_para(para):
    """过滤隐藏/保护段落"""
    if not para.text.strip():
        return False
    try:
        if para.font.hidden:
            return True
    except:
        pass
    return False

def set_complex_font(run, cn_font, en_font, size_pt, bold=False):
    """
    竞赛标准：中文/英文/数字 字体分离
    中文：cn_font | 英文/数字：en_font
    """
    # 中文字体
    run.font.name = cn_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    # 西文字体（数字+英文）
    run._element.rPr.rFonts.set(qn('w:ascii'), en_font)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), en_font)
    # 字号+加粗
    run.font.size = Pt(size_pt)
    run.font.bold = bold

def size_to_pt(size_str):
    """字号转磅值"""
    size_map = {
        "初号": 42, "小初": 36, "一号": 26, "小一": 24,
        "二号": 22, "小二": 18, "三号": 16, "小三": 15,
        "四号": 14, "小四": 12, "五号": 10.5, "小五": 9
    }
    return size_map.get(size_str, 12)

def process_doc(file, config, number_config, enable_title_regex,
                force_style, keep_spacing, clear_blank, max_blank):
    doc = Document(BytesIO(file.getvalue()))
    stats = {
        "一级标题":0,"二级标题":0,"三级标题":0,"正文":0,
        "表格":len(doc.tables),"图片":len([r for r in doc.element.xpath('.//a:blip')])
    }

    from core.title_recognizer import get_title_level

    # 处理正文段落（新增：上下文校验，避免正文列表误识别）
    prev_text = None  # 记录上一段文本，用于上下文校验
    for para in doc.paragraphs:
        if is_protected_para(para):
            continue
        # 传入上一段文本，进行上下文校验
        level = get_title_level(para.text, enable_title_regex, prev_text)
        stats[level] += 1
        # 更新上一段文本
        prev_text = para.text.strip()

        style = config[level]
        pt = size_to_pt(style["size"])
        # 固定西文字体为竞赛标准：Times New Roman
        en_font = "Times New Roman"
        cn_font = style["font"]

        # 逐段设置字体（中西分离）
        for run in para.runs:
            set_complex_font(run, cn_font, en_font, pt, style["bold"])

        # 首行缩进
        para.paragraph_format.first_line_indent = style["indent"] * 12700

    # 处理表格内文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style = config["表格"]
                    pt = size_to_pt(style["size"])
                    for run in para.runs:
                        set_complex_font(run, "宋体", "Times New Roman", pt, False)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, stats, 1.0, []